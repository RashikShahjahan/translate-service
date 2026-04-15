from pathlib import Path
from typing import Iterator

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_text_blocks(doc: Document) -> Iterator[str]:
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            text = Paragraph(child, doc).text.strip()
            if text:
                yield text
            continue

        if isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    yield "\t".join(cells)


def extract_text_from_docx(source_path: Path) -> str:
    doc = Document(source_path)
    return "\n\n".join(iter_text_blocks(doc))


def write_document_docx(document: dict, output_path: Path) -> None:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_paragraph(str(document["translated_text"] or ""))
    doc.save(output_path)
